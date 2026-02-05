"""
Phase 2: Semi-Structured Report Generation for CodeFlow Knowledge Graph

This script generates project status reports (DOCX format) that:
- Reference entities from entities.json
- 70% align with structured data (validates ground truth)
- 30% introduce controlled contradictions (tests entity resolution)
- Track all contradictions in metadata for later validation

Key Design Principles:
- Reports are realistic project status documents
- Mention co-occurrence ≠ semantic relationship
- Contradictions are intentional and tracked
- Each report cites provenance (who wrote it, when)
"""

import json
import yaml
import random
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


class ReportGenerator:
    """Generates semi-structured project status reports with controlled contradictions."""

    def __init__(self, config_path: str = "config.yaml"):
        """Initialize generator with configuration and entity data."""
        with open(config_path, 'r') as f:
            self.config = yaml.safe_load(f)

        # Set random seed
        random.seed(self.config['generation']['random_seed'])

        # Load entities and structured data
        self.entities = self._load_entities()
        self.employees_df = pd.read_csv(f"{self.config['paths']['structured_dir']}/employees.csv")
        self.projects_df = pd.read_csv(f"{self.config['paths']['structured_dir']}/projects.csv")
        self.products_df = pd.read_csv(f"{self.config['paths']['structured_dir']}/products.csv")
        self.policies_df = pd.read_csv(f"{self.config['paths']['structured_dir']}/policies.csv")
        self.assignments_df = pd.read_csv(f"{self.config['paths']['structured_dir']}/project_assignments.csv")

        # Track document metadata
        self.documents_metadata = []

        print("✓ Configuration and entity data loaded")

    def _load_entities(self) -> Dict:
        """Load canonical entity registry."""
        entities_path = f"{self.config['paths']['output_dir']}/entities.json"
        with open(entities_path, 'r') as f:
            return json.load(f)

    def _get_entity_by_id(self, entity_id: str) -> Dict:
        """Retrieve entity details by ID."""
        # Determine entity type from ID prefix
        if entity_id.startswith('emp_'):
            entities = self.entities['employees']
        elif entity_id.startswith('proj_'):
            entities = self.entities['projects']
        elif entity_id.startswith('prod_'):
            entities = self.entities['products']
        elif entity_id.startswith('pol_'):
            entities = self.entities['policies']
        elif entity_id.startswith('reg_'):
            entities = self.entities['regulations']
        else:
            return None

        for entity in entities:
            if entity['id'] == entity_id:
                return entity
        return None

    def _get_project_team(self, project_id: str) -> List[str]:
        """Get list of employee IDs assigned to a project."""
        assignments = self.assignments_df[
            self.assignments_df['project_id'] == project_id
            ]
        return assignments['employee_id'].tolist()

    def _get_random_quarter_year(self) -> Tuple[str, str]:
        """Generate random quarter and year for report timestamp."""
        quarters = ['Q1', 'Q2', 'Q3', 'Q4']
        years = ['2022', '2023', '2024']
        return random.choice(quarters), random.choice(years)

    def _select_contradiction_type(self) -> str:
        """Select type of contradiction to introduce."""
        types = self.config['contradictions']['types']
        return random.choice(types)

    def _introduce_project_assignment_contradiction(self,
                                                    project_id: str,
                                                    correct_team: List[str]) -> Tuple[str, Dict]:
        """
        Introduce a contradiction: mention employee on wrong project.

        Returns:
            Tuple of (employee_id_used, contradiction_metadata)
        """
        # Get an employee NOT on this project
        all_employees = [e['id'] for e in self.entities['employees']]
        wrong_employees = [e for e in all_employees if e not in correct_team]

        if not wrong_employees:
            return None, None

        wrong_emp_id = random.choice(wrong_employees)

        # Find what project they're actually on
        actual_assignment = self.assignments_df[
            self.assignments_df['employee_id'] == wrong_emp_id
            ]
        actual_project = actual_assignment['project_id'].iloc[0] if len(actual_assignment) > 0 else "none"

        contradiction = {
            'type': 'project_assignment',
            'entity': wrong_emp_id,
            'csv_value': actual_project,
            'document_value': project_id,
            'explanation': f"Report mentions {wrong_emp_id} on {project_id} but CSV shows {wrong_emp_id} on {actual_project}"
        }

        return wrong_emp_id, contradiction

    def _introduce_product_mention_contradiction(self) -> Tuple[str, Dict]:
        """
        Introduce a contradiction: mention product not in catalog.

        Returns:
            Tuple of (product_name, contradiction_metadata)
        """
        # Shadow IT products not in official catalog
        shadow_products = [
            "Microsoft Teams",
            "Notion",
            "Airtable",
            "Figma",
            "Miro",
            "Zoom",
            "Google Workspace"
        ]

        shadow_product = random.choice(shadow_products)

        contradiction = {
            'type': 'product_mention',
            'entity': shadow_product,
            'csv_value': 'not_in_catalog',
            'document_value': shadow_product,
            'explanation': f"Report mentions '{shadow_product}' which is not in products.csv"
        }

        return shadow_product, contradiction

    def _introduce_policy_reference_contradiction(self) -> Tuple[str, Dict]:
        """
        Introduce a contradiction: mention non-existent policy.

        Returns:
            Tuple of (policy_name, contradiction_metadata)
        """
        fake_policies = [
            "Remote Work Policy",
            "AI Usage Guidelines",
            "Social Media Policy",
            "BYOD Policy",
            "Incident Response Policy"
        ]

        fake_policy = random.choice(fake_policies)

        contradiction = {
            'type': 'policy_reference',
            'entity': fake_policy,
            'csv_value': 'not_in_catalog',
            'document_value': fake_policy,
            'explanation': f"Report mentions '{fake_policy}' which is not in policies.csv"
        }

        return fake_policy, contradiction

    def _format_name_for_report(self, employee: Dict, use_formal: bool = False) -> str:
        """
        Format employee name for report context.

        Args:
            employee: Employee entity dict
            use_formal: If True, use formal format (Ms./Mr. LastName)

        Returns:
            Formatted name string
        """
        if use_formal and random.random() < 0.3:
            # 30% chance of formal when requested
            prefix = random.choice(['Ms.', 'Mr.'])
            return f"{prefix} {employee['last_name']}"
        elif random.random() < 0.5:
            # 50% chance of first name only
            return employee['first_name']
        else:
            # Otherwise full name
            return employee['full_name']

    def _generate_report_content(self,
                                 project: Dict,
                                 introduce_contradictions: bool) -> Tuple[Dict, List[Dict], List[str]]:
        """
        Generate content for a project report.

        Args:
            project: Project entity dict
            introduce_contradictions: Whether to introduce contradictions

        Returns:
            Tuple of (content_dict, contradictions_list, entities_mentioned_list)
        """
        project_id = project['id']

        # Get actual team from CSV
        correct_team = self._get_project_team(project_id)

        # Select report author (lead from team)
        lead_assignment = self.assignments_df[
            (self.assignments_df['project_id'] == project_id) &
            (self.assignments_df['role'].str.contains('Lead|Manager', case=False, na=False))
            ]

        if len(lead_assignment) > 0:
            author_id = lead_assignment['employee_id'].iloc[0]
        else:
            author_id = random.choice(correct_team)

        author = self._get_entity_by_id(author_id)

        # Get team members (2-3 to mention)
        team_to_mention = random.sample(correct_team, min(3, len(correct_team)))

        # Get products used (1-2 from catalog)
        products_to_use = random.sample(
            [p['id'] for p in self.entities['products']],
            random.randint(1, 2)
        )

        # Get policies referenced (0-1)
        policies_to_reference = []
        if random.random() < 0.4:  # 40% chance of mentioning policy
            policies_to_reference = random.sample(
                [p['id'] for p in self.entities['policies']],
                1
            )

        # Track entities mentioned
        entities_mentioned = [project_id, author_id] + team_to_mention + products_to_use + policies_to_reference

        # Track contradictions
        contradictions = []

        # CONTRADICTION LOGIC: 30-40% of reports have contradictions
        if introduce_contradictions and random.random() < self.config['contradictions']['target_percentage']:
            num_contradictions = random.randint(1, 2)

            for _ in range(num_contradictions):
                contradiction_type = self._select_contradiction_type()

                if contradiction_type == 'project_assignment':
                    wrong_emp_id, contradiction = self._introduce_project_assignment_contradiction(
                        project_id, correct_team
                    )
                    if wrong_emp_id:
                        team_to_mention.append(wrong_emp_id)
                        entities_mentioned.append(wrong_emp_id)
                        contradictions.append(contradiction)

                elif contradiction_type == 'product_mention':
                    shadow_product, contradiction = self._introduce_product_mention_contradiction()
                    products_to_use.append(shadow_product)  # This is a string, not ID
                    contradictions.append(contradiction)

                elif contradiction_type == 'policy_reference':
                    fake_policy, contradiction = self._introduce_policy_reference_contradiction()
                    policies_to_reference.append(fake_policy)  # This is a string, not ID
                    contradictions.append(contradiction)

        # Generate quarter/year
        quarter, year = self._get_random_quarter_year()

        # Build content
        content = {
            'project': project,
            'author': author,
            'quarter': quarter,
            'year': year,
            'team_mentions': [self._get_entity_by_id(emp_id) if emp_id.startswith('emp_') else None
                              for emp_id in team_to_mention],
            'products': [(self._get_entity_by_id(p_id) if p_id.startswith('prod_') else {'name': p_id})
                         for p_id in products_to_use],
            'policies': [(self._get_entity_by_id(p_id) if p_id.startswith('pol_') else {'name': p_id})
                         for p_id in policies_to_reference]
        }

        return content, contradictions, entities_mentioned

    def _create_docx_report(self, content: Dict, doc_id: str) -> str:
        """
        Create DOCX file for report.

        Args:
            content: Report content dict
            doc_id: Document ID (e.g., doc_001)

        Returns:
            Filename of created document
        """
        doc = Document()

        # Title
        title = doc.add_heading(
            f"Project {content['project']['name']} - {content['quarter']} {content['year']} Status Report",
            0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Author and metadata
        meta = doc.add_paragraph()
        meta.add_run('Lead: ').bold = True
        meta.add_run(f"{content['author']['full_name']} ({content['author']['role']})\n")
        meta.add_run('Status: ').bold = True

        # Determine status display
        status = content['project']['status']
        if status == 'Active':
            status_text = 'On track'
        elif status == 'Completed':
            status_text = 'Completed'
        else:
            status_text = 'Planned'
        meta.add_run(status_text)

        doc.add_paragraph()  # Spacing

        # Summary section
        doc.add_heading('Summary', level=1)

        # Paragraph 1: Project overview
        p1 = doc.add_paragraph()
        p1.add_run(f"Project {content['project']['name']} ")

        if status == 'Completed':
            p1.add_run("has been successfully completed. ")
        elif status == 'Active':
            p1.add_run("continues to make progress this quarter. ")
        else:
            p1.add_run("is in the planning phase. ")

        p1.add_run(content['project']['description'])

        # Paragraph 2: Team and tools
        p2 = doc.add_paragraph()

        # Mention team members
        team_names = []
        for emp in content['team_mentions']:
            if emp:
                team_names.append(self._format_name_for_report(emp))

        if len(team_names) == 1:
            p2.add_run(f"{team_names[0]} has been leading the effort. ")
        elif len(team_names) == 2:
            p2.add_run(f"{team_names[0]} and {team_names[1]} have been collaborating closely. ")
        else:
            p2.add_run(f"{', '.join(team_names[:-1])}, and {team_names[-1]} have been working together. ")

        # Mention products
        if content['products']:
            product_names = [p['name'] for p in content['products']]
            if len(product_names) == 1:
                p2.add_run(f"The team is using {product_names[0]} ")
            else:
                p2.add_run(f"The team is leveraging {', '.join(product_names[:-1])} and {product_names[-1]} ")
            p2.add_run("to support the project's objectives.")

        # Paragraph 3: Compliance/policies (if applicable)
        if content['policies']:
            p3 = doc.add_paragraph()
            policy_names = [p['name'] for p in content['policies']]
            p3.add_run(f"The project maintains alignment with {policy_names[0]} ")
            p3.add_run("to ensure compliance and risk management standards are met.")

        doc.add_paragraph()  # Spacing

        # Blockers/Risks section
        doc.add_heading('Blockers/Risks', level=1)

        risks = [
            "Resource allocation may need adjustment for Q4 deliverables",
            "Integration dependencies with external systems require coordination",
            "Timeline is tight for the upcoming compliance review",
            "Vendor response times have been slower than expected",
            "Technical debt in legacy systems may impact migration schedule"
        ]

        doc.add_paragraph(random.choice(risks), style='List Bullet')
        if random.random() < 0.5:
            doc.add_paragraph(random.choice([r for r in risks]), style='List Bullet')

        doc.add_paragraph()  # Spacing

        # Next Steps section
        doc.add_heading('Next Steps', level=1)

        next_steps = [
            "Finalize architecture design with stakeholder review",
            "Complete security audit and penetration testing",
            "Schedule training sessions for end users",
            "Coordinate with legal team on data processing agreements",
            "Deploy to staging environment for UAT",
            "Update project documentation and runbooks"
        ]

        doc.add_paragraph(random.choice(next_steps), style='List Bullet')
        doc.add_paragraph(random.choice([s for s in next_steps]), style='List Bullet')

        # Save document
        filename = f"project_report_{doc_id.split('_')[1]}.docx"
        filepath = f"{self.config['paths']['semi_structured_dir']}/{filename}"
        doc.save(filepath)

        return filename

    def generate_reports(self, num_reports: int = 6) -> List[Dict]:
        """
        Generate all project reports.

        Args:
            num_reports: Number of reports to generate (default 4-6)

        Returns:
            List of document metadata dicts
        """
        print(f"\n[Phase 2] Generating {num_reports} project status reports...")

        # Select projects to write reports about (can repeat projects)
        available_projects = [p for p in self.entities['projects']]
        selected_projects = random.sample(available_projects, min(num_reports, len(available_projects)))

        # If we need more reports than projects, repeat some
        while len(selected_projects) < num_reports:
            selected_projects.append(random.choice(available_projects))

        documents_metadata = []

        for i, project in enumerate(selected_projects, 1):
            doc_id = f"doc_{i:03d}"

            # Determine if this report should have contradictions
            # Target: 30-40% with contradictions, but vary it
            introduce_contradictions = (i % 3 == 0)  # Every 3rd report

            # Generate content
            content, contradictions, entities_mentioned = self._generate_report_content(
                project,
                introduce_contradictions
            )

            # Create DOCX
            filename = self._create_docx_report(content, doc_id)

            # Calculate alignment confidence
            if contradictions:
                confidence_alignment = 1.0 - (len(contradictions) * 0.15)  # Each contradiction reduces confidence
            else:
                confidence_alignment = 1.0

            # Build metadata
            metadata = {
                'id': doc_id,
                'filename': filename,
                'type': 'semi_structured',
                'source': 'generated',
                'timestamp': f"{content['year']}-{content['quarter'].replace('Q', '0')[:2]}",
                'entities_mentioned': [
                    {'id': ent_id, 'mention_text': self._get_entity_by_id(ent_id)['name'] if self._get_entity_by_id(
                        ent_id) and 'name' in self._get_entity_by_id(ent_id) else ent_id}
                    for ent_id in entities_mentioned if ent_id.startswith(('emp_', 'proj_', 'prod_', 'pol_'))
                ],
                'contradictions': contradictions,
                'confidence_alignment': round(confidence_alignment, 2)
            }

            documents_metadata.append(metadata)

            status_icon = "⚠️" if contradictions else "✓"
            print(f"  {status_icon} Generated {filename} ({len(contradictions)} contradictions)")

        return documents_metadata

    def save_metadata(self, documents_metadata: List[Dict]):
        """Save document metadata to JSON."""
        metadata_path = f"{self.config['paths']['output_dir']}/reports_metadata.json"

        with open(metadata_path, 'w') as f:
            json.dump({
                'generation_timestamp': datetime.now().isoformat(),
                'total_reports': len(documents_metadata),
                'documents': documents_metadata
            }, f, indent=2)

        print(f"\n✓ Saved metadata to reports_metadata.json")

    def print_summary(self, documents_metadata: List[Dict]):
        """Print generation summary statistics."""
        total = len(documents_metadata)
        with_contradictions = sum(1 for d in documents_metadata if d['contradictions'])
        total_contradictions = sum(len(d['contradictions']) for d in documents_metadata)

        print(f"\n{'=' * 60}")
        print("Phase 2 Summary:")
        print(f"  Total reports generated: {total}")
        print(f"  Reports with contradictions: {with_contradictions} ({with_contradictions / total * 100:.1f}%)")
        print(f"  Total contradictions: {total_contradictions}")
        print(f"  Average contradictions per report: {total_contradictions / total:.2f}")
        print(f"{'=' * 60}")

    def run(self, num_reports: int = 6):
        """Execute complete report generation pipeline."""
        print("=" * 60)
        print("CodeFlow Knowledge Graph - Phase 2: Report Generation")
        print("=" * 60)

        # Generate reports
        documents_metadata = self.generate_reports(num_reports)

        # Save metadata
        self.save_metadata(documents_metadata)

        # Print summary
        self.print_summary(documents_metadata)

        print("\n✅ Phase 2 complete!")
        print("\nNext Steps:")
        print("  1. Review generated reports in data/semi_structured/")
        print("  2. Examine reports_metadata.json for contradiction tracking")
        print("  3. Proceed to Phase 3: generate_emails.py")


def main():
    """Main entry point."""
    generator = ReportGenerator("config.yaml")
    generator.run(num_reports=6)


if __name__ == "__main__":
    main()